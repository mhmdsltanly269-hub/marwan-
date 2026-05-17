import os
import re
import io
import json
from datetime import datetime, timedelta
from functools import wraps

from flask import (Flask, render_template, request, jsonify, 
                   session, redirect, url_for, send_file)
from flask_cors import CORS
from flask_bcrypt import Bcrypt
from flask_session import Session
from flask_talisman import Talisman
from dotenv import load_dotenv
import openai
from pymongo import MongoClient
from bson.objectid import ObjectId
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import mammoth  # لتحويل .docx إلى نص
from werkzeug.utils import secure_filename
import cloudinary
import cloudinary.uploader
import cloudinary.api
from apscheduler.schedulers.background import BackgroundScheduler
import pytz
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

# تحميل المتغيرات البيئية
load_dotenv()

# إعداد التطبيق
app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'mervan-academic-secret-2024')
app.config['SESSION_TYPE'] = 'filesystem'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=24)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50 ميجابايت

# إعدادات الأمان
Talisman(app, content_security_policy=None)
Session(app)
CORS(app)
bcrypt = Bcrypt(app)
limiter = Limiter(key_func=get_remote_address, app=app, default_limits=["200 per day", "50 per hour"])

# إعدادات Cloudinary للصور (للأفاتار والملفات)
cloudinary.config(
    cloud_name=os.getenv('CLOUDINARY_CLOUD_NAME'),
    api_key=os.getenv('CLOUDINARY_API_KEY'),
    api_secret=os.getenv('CLOUDINARY_API_SECRET')
)

# إعدادات OpenAI (أو يمكن استخدام Claude)
openai.api_key = os.getenv('OPENAI_API_KEY')

# إعدادات MongoDB
client = MongoClient(os.getenv('MONGODB_URI', 'mongodb://localhost:27017/'))
db = client['mervan_academic']
users_collection = db['users']
texts_collection = db['texts']

# ============ نظام التوثيق ============
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login_page'))
        return f(*args, **kwargs)
    return decorated_function

# ============ صفحات الواجهة ============
@app.route('/')
def login_page():
    return render_template('index.html')

@app.route('/dashboard')
@login_required
def dashboard():
    user = users_collection.find_one({'_id': ObjectId(session['user_id'])})
    return render_template('dashboard.html', user=user)

# ============ API: تسجيل الدخول والتسجيل ============
@app.route('/api/register', methods=['POST'])
@limiter.limit("10 per hour")
def register():
    data = request.json
    username = data.get('username')
    email = data.get('email')
    password = data.get('password')
    
    if not all([username, email, password]):
        return jsonify({'success': False, 'message': 'جميع الحقول مطلوبة'}), 400
    
    if users_collection.find_one({'email': email}):
        return jsonify({'success': False, 'message': 'البريد الإلكتروني مسجل مسبقاً'}), 400
    
    hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
    user = {
        'username': username,
        'email': email,
        'password': hashed_password,
        'created_at': datetime.utcnow(),
        'usage_count': 0,
        'plan': 'free'
    }
    result = users_collection.insert_one(user)
    session['user_id'] = str(result.inserted_id)
    
    return jsonify({'success': True, 'message': 'تم التسجيل بنجاح'})

@app.route('/api/login', methods=['POST'])
@limiter.limit("20 per hour")
def login():
    data = request.json
    email = data.get('email')
    password = data.get('password')
    
    user = users_collection.find_one({'email': email})
    if user and bcrypt.check_password_hash(user['password'], password):
        session['user_id'] = str(user['_id'])
        return jsonify({'success': True, 'message': 'مرحباً بك في مروان أكاديمي'})
    
    return jsonify({'success': False, 'message': 'بيانات الدخول غير صحيحة'}), 401

@app.route('/api/logout')
@login_required
def logout():
    session.clear()
    return jsonify({'success': True})

# ============ API: معالجة النصوص ============
@app.route('/api/rephrase', methods=['POST'])
@login_required
@limiter.limit("30 per hour")
def rephrase_text():
    data = request.json
    text = data.get('text', '')
    style = data.get('style', 'academic')  # academic, creative, summary
    
    if not text:
        return jsonify({'success': False, 'message': 'النص مطلوب'}), 400
    
    word_count = len(text.split())
    if word_count > 20000:
        return jsonify({'success': False, 'message': 'الحد الأقصى 20,000 كلمة'}), 400
    
    try:
        # إعداد الموجه (Prompt) حسب النمط
        prompts = {
            'academic': """أنت خبير لغوي أكاديمي رفيع المستوى. مهمتك إعادة صياغة النص التالي بأسلوب بشري أكاديمي محترف تماماً، مع مراعاة:
1. إزالة أي طابع آلي أو ركاكة الذكاء الاصطناعي
2. استخدام مفردات أكاديمية دقيقة ومتنوعة
3. الحفاظ على التسلسل المنطقي والاستشهادات
4. تحسين بنية الجمل لتكون سلسة واحترافية
5. الحفاظ على المعنى الأصلي بالكامل

النص المراد صياغته:
{text}""",
            'creative': """أنت كاتب إبداعي محترف. أعد صياغة النص التالي بأسلوب إبداعي جذاب، مع الحفاظ على المعنى وإضافة لمسات بلاغية تجعله أكثر حيوية وإنسانية:
{text}""",
            'summary': """أنت خبير في التلخيص الأكاديمي. لخص النص التالي بأسلوب بشري احترافي، مع الحفاظ على النقاط الرئيسية وتقديمها بشكل متماسك:
{text}"""
        }
        
        prompt = prompts.get(style, prompts['academic']).format(text=text)
        
        response = openai.ChatCompletion.create(
            model="gpt-4o",  # أفضل نموذج للغة العربية والأكاديمية
            messages=[
                {"role": "system", "content": "أنت مساعد أكاديمي خبير في صياغة النصوص العربية بأسلوب بشري احترافي."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=4000
        )
        
        rephrased_text = response.choices[0].message['content']
        
        # حفظ في قاعدة البيانات
        text_record = {
            'user_id': session['user_id'],
            'original_text': text,
            'rephrased_text': rephrased_text,
            'style': style,
            'word_count': word_count,
            'created_at': datetime.utcnow()
        }
        texts_collection.insert_one(text_record)
        
        # تحديث عداد الاستخدام
        users_collection.update_one(
            {'_id': ObjectId(session['user_id'])},
            {'$inc': {'usage_count': 1}}
        )
        
        return jsonify({
            'success': True,
            'rephrased_text': rephrased_text,
            'word_count': word_count
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

# ============ API: رفع الملفات ومعالجتها ============
@app.route('/api/upload-file', methods=['POST'])
@login_required
@limiter.limit("20 per hour")
def upload_file():
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'لم يتم رفع ملف'}), 400
    
    file = request.files['file']
    style = request.form.get('style', 'academic')
    
    if file.filename == '':
        return jsonify({'success': False, 'message': 'الملف فارغ'}), 400
    
    # السماح بأنواع الملفات
    allowed_extensions = {'txt', 'pdf', 'docx', 'doc'}
    file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
    
    if file_ext not in allowed_extensions:
        return jsonify({'success': False, 'message': 'نوع الملف غير مدعوم'}), 400
    
    try:
        # استخراج النص حسب نوع الملف
        text = ""
        if file_ext == 'txt':
            text = file.read().decode('utf-8')
        elif file_ext == 'pdf':
            pdf_reader = PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text()
        elif file_ext in ['docx', 'doc']:
            result = mammoth.extract_raw_text(file)
            text = result.value
        
        if not text.strip():
            return jsonify({'success': False, 'message': 'الملف فارغ أو غير قابل للقراءة'}), 400
        
        # إعادة الصياغة
        prompt = """أنت خبير لغوي أكاديمي. أعد صياغة النص التالي من ملف {file_name} بأسلوب بشري أكاديمي:
{text}""".format(file_name=file.filename, text=text)
        
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "أنت محرر أكاديمي محترف."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=4000
        )
        
        rephrased_text = response.choices[0].message['content']
        
        return jsonify({
            'success': True,
            'original_text': text[:500] + '...' if len(text) > 500 else text,
            'rephrased_text': rephrased_text,
            'file_name': file.filename
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

# ============ API: تحميل النتيجة كملف وورد ============
@app.route('/api/download-word', methods=['POST'])
@login_required
def download_word():
    data = request.json
    text = data.get('text', '')
    title = data.get('title', 'مروان أكاديمي - نص معاد صياغته')
    
    if not text:
        return jsonify({'success': False, 'message': 'النص مطلوب'}), 400
    
    try:
        # إنشاء مستند وورد احترافي
        doc = Document()
        
        # إعدادات الصفحة
        section = doc.sections[0]
        section.page_width = Inches(8.27)  # A4
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        # إضافة رأس الصفحة
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_paragraph.add_run('🎓 مروان أكاديمي - Mervan Academic')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(212, 175, 55)  # ذهبي
        
        # عنوان المستند
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_paragraph.add_run(title)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(13, 27, 42)  # كحلي
        
        # خط فاصل ذهبي
        separator = doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = separator.add_run('━' * 50)
        run.font.color.rgb = RGBColor(212, 175, 55)
        run.font.size = Pt(8)
        
        # المحتوى
        paragraphs = text.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(para_text)
                run.font.size = Pt(14)
                run.font.name = 'Traditional Arabic'
        
        # تذييل الصفحة
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_paragraph.add_run(f'© 2024 مروان أكاديمي | تمت الصياغة: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        # حفظ الملف في الذاكرة
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'مروان_اكاديمي_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        )
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

# ============ API: السجل ============
@app.route('/api/history')
@login_required
def get_history():
    records = texts_collection.find({'user_id': session['user_id']}).sort('created_at', -1).limit(20)
    history = []
    for record in records:
        history.append({
            'id': str(record['_id']),
            'original_preview': record['original_text'][:100] + '...',
            'rephrased_preview': record['rephrased_text'][:100] + '...',
            'style': record['style'],
            'word_count': record['word_count'],
            'created_at': record['created_at'].strftime('%Y-%m-%d %H:%M')
        })
    return jsonify({'success': True, 'history': history})

# ============ تشغيل التطبيق ============
if __name__ == '__main__':
    app.run(debug=False, host='0.0.0.0', port=5000)